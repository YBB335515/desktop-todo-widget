"""
学习笔记生成器 - 由 /老师 skill 调用。
分析 Python 代码文件，生成包含代码清单、白话解释、API教程和术语表的 Word 文档。

用法:
    python generate_notes.py <target_file>

输出:
    <target_file>_学习笔记.docx
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor


def build_doc(target_path):
    """Build the learning notes Word document for a target Python file."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    filename = os.path.basename(target_path)
    with open(target_path, "r", encoding="utf-8") as f:
        code_lines = f.readlines()

    # ============================================================
    # 封面 / 标题
    # ============================================================
    title = doc.add_heading(f'{filename} — 学习笔记', level=0)
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
    p = doc.add_paragraph()
    run = p.add_run(f'源文件: {target_path}\n总行数: {len(code_lines)}')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)

    # ============================================================
    # 第一章: 代码清单
    # ============================================================
    doc.add_heading('第一章 代码清单', level=1)
    for i, line in enumerate(code_lines, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # Line number
        run = p.add_run(f"{i:4d}  ")
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        # Code
        run = p.add_run(line.rstrip())
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Add a page break after code listing
    doc.add_page_break()

    # ============================================================
    # 第二章: 逐段白话解释 (placeholder - AI fills this)
    # ============================================================
    doc.add_heading('第二章 逐段白话解释', level=1)
    doc.add_paragraph('（此章节由 Claude AI 分析代码后填充）')

    # ============================================================
    # 第三章: API 使用教程 (placeholder - AI fills this)
    # ============================================================
    doc.add_heading('第三章 API 使用教程', level=1)
    doc.add_paragraph('（此章节由 Claude AI 分析代码后填充）')

    # ============================================================
    # 第四章: 术语表 (placeholder - AI fills this)
    # ============================================================
    doc.add_heading('第四章 术语表', level=1)
    doc.add_paragraph('（此章节由 Claude AI 分析代码后填充）')

    # ============================================================
    # Save
    # ============================================================
    base = os.path.splitext(target_path)[0]
    output = f"{base}_学习笔记.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python generate_notes.py <target_file>")
        sys.exit(1)
    target = sys.argv[1]
    if not os.path.isfile(target):
        print(f"文件不存在: {target}")
        sys.exit(1)
    out = build_doc(target)
    print(out)
